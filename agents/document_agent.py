"""
Document Agent for SupportGenie
Handles multi-format document upload and processing
"""

import PyPDF2
import os
import json
import csv
import xml.etree.ElementTree as ET
from typing import Dict
from services.gemini_service import GeminiService
from services.database import Database

# Import optional dependencies
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False


class DocumentAgent:
    """Agent responsible for processing and managing documents in multiple formats"""
    
    def __init__(self, gemini_service: GeminiService, database: Database):
        """Initialize Document Agent with required services"""
        self.gemini = gemini_service
        self.db = database
    
    def process_document(self, file_path: str) -> Dict[str, any]:
        """
        Process a document file: extract text, structure with AI, save to database
        Supports: PDF, TXT, DOCX, XLSX, XLS, CSV, JSON, MD, XML
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict with success status, document_id, and filename
        """
        try:
            # Extract filename and extension
            filename = os.path.basename(file_path)
            file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            
            # Route to appropriate extractor based on file type
            if file_extension == 'pdf':
                original_content = self._extract_text_from_pdf(file_path)
            elif file_extension == 'txt':
                original_content = self._extract_text_from_txt(file_path)
            elif file_extension == 'docx':
                original_content = self._extract_text_from_docx(file_path)
            elif file_extension == 'xlsx':
                original_content = self._extract_text_from_xlsx(file_path)
            elif file_extension == 'xls':
                original_content = self._extract_text_from_xls(file_path)
            elif file_extension == 'csv':
                original_content = self._extract_text_from_csv(file_path)
            elif file_extension == 'json':
                original_content = self._extract_text_from_json(file_path)
            elif file_extension == 'md':
                original_content = self._extract_text_from_markdown(file_path)
            elif file_extension == 'xml':
                original_content = self._extract_text_from_xml(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_extension}'
                }
            
            if not original_content or len(original_content.strip()) < 10:
                return {
                    'success': False,
                    'error': 'File appears to be empty or has insufficient content'
                }
            
            # Process with Gemini to structure the content
            structured_content = self.gemini.process_document(original_content)
            
            # Save to database
            document_id = self.db.save_document(
                filename=filename,
                original_content=original_content,
                structured_content=structured_content
            )
            
            return {
                'success': True,
                'document_id': document_id,
                'filename': filename,
                'content_length': len(original_content)
            }
            
        except FileNotFoundError:
            return {
                'success': False,
                'error': 'File not found'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text content from PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text_content = ""
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n\n"
            
            return text_content.strip()
            
        except PyPDF2.errors.PdfReadError as e:
            raise Exception(f"Corrupted or invalid PDF file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_text_from_txt(self, txt_path: str) -> str:
        """
        Extract text content from TXT file
        
        Args:
            txt_path: Path to the TXT file
            
        Returns:
            Extracted text content
        """
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text content from DOCX file
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(docx_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _extract_text_from_csv(self, csv_path: str) -> str:
        """
        Extract text content from CSV file
        
        Args:
            csv_path: Path to the CSV file
            
        Returns:
            Extracted text content formatted as readable text
        """
        try:
            text_content = []
            
            with open(csv_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                
                # Get headers
                headers = next(csv_reader, None)
                if headers:
                    text_content.append("Headers: " + " | ".join(headers))
                    text_content.append("-" * 50)
                
                # Get rows
                for row in csv_reader:
                    if row:
                        text_content.append(" | ".join(row))
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading CSV file: {str(e)}")
    
    def _extract_text_from_json(self, json_path: str) -> str:
        """
        Extract text content from JSON file
        
        Args:
            json_path: Path to the JSON file
            
        Returns:
            Extracted text content formatted as readable text
        """
        try:
            with open(json_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)
            
        except json.JSONDecodeError as e:
            raise Exception(f"Invalid JSON file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading JSON file: {str(e)}")
    
    def _extract_text_from_markdown(self, md_path: str) -> str:
        """
        Extract text content from Markdown file
        
        Args:
            md_path: Path to the Markdown file
            
        Returns:
            Extracted text content (preserving markdown formatting)
        """
        try:
            with open(md_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading Markdown file: {str(e)}")
    
    def _extract_text_from_xml(self, xml_path: str) -> str:
        """
        Extract text content from XML file
        
        Args:
            xml_path: Path to the XML file
            
        Returns:
            Extracted text content formatted as readable text
        """
        try:
            tree = ET.parse(xml_path)
            root = tree.getroot()
            
            def extract_from_element(element, level=0):
                """Recursively extract text from XML elements"""
                text_parts = []
                indent = "  " * level
                
                # Add element tag and attributes
                if element.attrib:
                    attrs = ", ".join([f"{k}={v}" for k, v in element.attrib.items()])
                    text_parts.append(f"{indent}{element.tag} ({attrs}):")
                else:
                    text_parts.append(f"{indent}{element.tag}:")
                
                # Add element text
                if element.text and element.text.strip():
                    text_parts.append(f"{indent}  {element.text.strip()}")
                
                # Process child elements
                for child in element:
                    text_parts.extend(extract_from_element(child, level + 1))
                
                return text_parts
            
            all_text = extract_from_element(root)
            return "\n".join(all_text)
            
        except ET.ParseError as e:
            raise Exception(f"Invalid XML file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading XML file: {str(e)}")
    
    def _extract_text_from_xlsx(self, xlsx_path: str) -> str:
        """
        Extract text content from XLSX file (Excel 2007+)
        
        Args:
            xlsx_path: Path to the XLSX file
            
        Returns:
            Extracted text content formatted as readable text
        """
        if not OPENPYXL_AVAILABLE:
            raise Exception("openpyxl library not installed. Install with: pip install openpyxl")
        
        try:
            workbook = load_workbook(xlsx_path, data_only=True)
            text_content = []
            
            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"=== Sheet: {sheet_name} ===")
                text_content.append("")
                
                # Get all rows
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # Only add non-empty rows
                        text_content.append(" | ".join(row_values))
                
                text_content.append("")  # Add blank line between sheets
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading XLSX file: {str(e)}")
    
    def _extract_text_from_xls(self, xls_path: str) -> str:
        """
        Extract text content from XLS file (Excel 97-2003)
        
        Args:
            xls_path: Path to the XLS file
            
        Returns:
            Extracted text content formatted as readable text
        """
        if not XLRD_AVAILABLE:
            raise Exception("xlrd library not installed. Install with: pip install xlrd")
        
        try:
            workbook = xlrd.open_workbook(xls_path)
            text_content = []
            
            # Process each sheet
            for sheet_index in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_index)
                text_content.append(f"=== Sheet: {sheet.name} ===")
                text_content.append("")
                
                # Get all rows
                for row_index in range(sheet.nrows):
                    row_values = []
                    for col_index in range(sheet.ncols):
                        cell = sheet.cell(row_index, col_index)
                        row_values.append(str(cell.value) if cell.value else "")
                    
                    if any(row_values):  # Only add non-empty rows
                        text_content.append(" | ".join(row_values))
                
                text_content.append("")  # Add blank line between sheets
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading XLS file: {str(e)}")
    
    def get_all_documents(self) -> list:
        """Get all uploaded documents"""
        return self.db.get_all_documents()
    
    def get_document_content(self, document_id: int) -> Dict:
        """Get specific document content"""
        return self.db.get_document_by_id(document_id)

